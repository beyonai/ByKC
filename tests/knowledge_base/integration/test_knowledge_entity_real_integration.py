"""KnowledgeEntity API integration tests with deterministic model doubles.

The application, database, object storage, and Redis use the production wiring.
LLM and embedding boundaries are replaced at the client methods so these tests
remain deterministic and cannot send document or entity content externally.
"""

from __future__ import annotations

import asyncio
import concurrent.futures
import hashlib
import io
import json
import re
import time
from collections.abc import Iterable
from contextlib import contextmanager
from decimal import Decimal
from pathlib import Path
from typing import Any
from uuid import uuid4

import pytest
from by_framework.common.redis_client import close_redis
from fastapi.testclient import TestClient

import by_qa.main as main_module
from by_qa.config import get_settings
from by_qa.knowledge_base.api.knowledge_entity_schemas import EntityDiscoveryRequest
from by_qa.knowledge_base.infrastructure.database import build_connection_factory
from by_qa.knowledge_base.repositories.knowledge_entity_repository import (
    KnowledgeEntityRepository,
)
from by_qa.knowledge_base.services.bootstrap_service import split_sql_statements
from by_qa.knowledge_base.services.embedding_query_service import EmbeddingQueryService
from by_qa.knowledge_base.services.knowledge_entity_intelligence import (
    OpenAICompatibleKnowledgeEntityLLM,
)
from by_qa.knowledge_base.services.markdown_update_summary_service import (
    MarkdownUpdateSummaryService,
)
from by_qa.knowledge_build.services.document_chunking_service import (
    DocumentChunkingService,
)

TERMINAL_TASK_STATUSES = frozenset({"SUCCEEDED", "FAILED", "CANCELLED", "SKIPPED"})
TASK_TIMEOUT_SECONDS = 300.0


def _mock_entity_name(markdown: str) -> tuple[str, str]:
    named_entity = re.search(r"(?m)^([^\n.]+?)\s+is also called\s+[^\n.]+\.", markdown)
    if named_entity is not None:
        return named_entity.group(1).strip(), named_entity.group(0).strip()
    entity_link = re.search(
        r"(?m)^([^\n]*\[([^\]]+)\]\([^)]+\)[^\n]*)$",
        markdown,
    )
    if entity_link is not None:
        return entity_link.group(2).strip(), entity_link.group(1).strip()
    heading = re.search(r"(?m)^\s*#\s+(.+?)\s*$", markdown)
    if heading is not None:
        return heading.group(1).strip(), heading.group(0).strip()
    for line in markdown.splitlines():
        evidence = line.strip()
        if evidence:
            return evidence.lstrip("# ").strip(), evidence
    raise AssertionError("entity discovery mock received an empty document")


async def _mock_llm_cache_identity(
    llm: OpenAICompatibleKnowledgeEntityLLM,
) -> str:
    del llm
    return "knowledge-entity-integration-mock-v1"


async def _mock_llm_complete(
    llm: OpenAICompatibleKnowledgeEntityLLM,
    messages,
    *,
    json_mode: bool = False,
) -> str:
    del llm, json_mode
    system_prompt = str(messages[0]["content"])
    user_prompt = str(messages[-1]["content"])

    if "知识实体同义词裁决器" in system_prompt:
        return json.dumps(
            {
                "decision": "DIFFERENT",
                "selectedCandidateId": None,
                "canonicalName": None,
                "aliasToAdd": None,
                "reasonCode": "MOCK_DIFFERENT",
            },
            ensure_ascii=False,
        )

    if "核心对象实体发现器" in system_prompt:
        entity_name, evidence = _mock_entity_name(user_prompt)
        return json.dumps(
            [
                {
                    "entityName": entity_name,
                    "subjectEntityName": "",
                    "localName": entity_name,
                    "identityScope": "global",
                    "isEvent": False,
                    "evidence": evidence,
                    "aliases": [],
                }
            ],
            ensure_ascii=False,
        )

    if "updating exactly one existing KnowledgeEntity" in system_prompt:
        identity = re.search(r"(?m)^- entityName:\s*(.+?)\s*$", user_prompt)
        assert identity is not None, "enrichment mock did not receive entityName"
        source_reference = re.search(
            r"Source reference[^:]*:\s*(\[[^\n]+\]\([^\n]+\))", user_prompt
        )
        assert source_reference is not None, (
            "enrichment mock did not receive a Markdown source reference"
        )
        entity_name = identity.group(1).strip()
        reference = source_reference.group(1)
        digest = hashlib.sha256(user_prompt.encode("utf-8")).hexdigest()[:12]
        existing = re.search(
            r"Existing Markdown .*?:\n(.*?)\n\nSoft template guidance",
            user_prompt,
            flags=re.DOTALL,
        )
        assert existing is not None, "enrichment mock did not receive existing Markdown"
        existing_markdown = re.sub(
            r"\A---[ \t]*\n.*?\n---[ \t]*(?:\n|\Z)",
            "",
            existing.group(1).strip(),
            count=1,
            flags=re.DOTALL,
        ).strip()
        update_section = (
            f"## 集成测试更新\n\n已根据当前关联证据更新 {digest}。{reference}"
        )
        existing_markdown = re.sub(
            r"(?ms)^## 集成测试更新\s*$.*?(?=^#{1,2}\s|\Z)",
            "",
            existing_markdown,
        ).rstrip()
        await asyncio.sleep(0.05)
        markdown = f"{existing_markdown}\n\n{update_section}\n"
        if not markdown.lstrip().startswith(f"# {entity_name}"):
            markdown = f"# {entity_name}\n\n{markdown}"
        return json.dumps(
            {"markdown": markdown, "relations": [], "warnings": []},
            ensure_ascii=False,
        )

    raise AssertionError(
        "unexpected KnowledgeEntity LLM prompt; add an explicit deterministic mock"
    )


def _mock_embedding_vector(text: str, dimension: int) -> list[float]:
    del text
    assert dimension > 0
    vector = [0.0] * dimension
    vector[0] = 1.0
    return vector


async def _mock_query_embedding(
    embedding_service: EmbeddingQueryService, query: str
) -> list[float]:
    del embedding_service
    return _mock_embedding_vector(query, get_settings().embedding_dimension)


def _mock_chunk_embeddings(
    self: DocumentChunkingService, texts: list[str]
) -> list[list[float]]:
    return [_mock_embedding_vector(text, self.embedding_dimension) for text in texts]


async def _mock_update_summary(
    summary_service: MarkdownUpdateSummaryService,
    old_markdown: str,
    new_markdown: str,
) -> None:
    del summary_service, old_markdown, new_markdown
    return None


@pytest.fixture(autouse=True)
def _mock_external_model_services(monkeypatch: pytest.MonkeyPatch) -> None:
    """Keep integration coverage while forbidding external model requests."""

    monkeypatch.setattr(
        OpenAICompatibleKnowledgeEntityLLM,
        "cache_identity",
        _mock_llm_cache_identity,
    )
    monkeypatch.setattr(
        OpenAICompatibleKnowledgeEntityLLM,
        "complete",
        _mock_llm_complete,
    )
    monkeypatch.setattr(EmbeddingQueryService, "embed_query", _mock_query_embedding)
    monkeypatch.setattr(
        DocumentChunkingService,
        "_request_embeddings",
        _mock_chunk_embeddings,
    )
    monkeypatch.setattr(
        MarkdownUpdateSummaryService,
        "generate_llm_summary",
        _mock_update_summary,
    )


@contextmanager
def _real_test_client() -> Iterable[TestClient]:
    """Run one real lifespan and release Redis on that lifespan's event loop."""

    # Other integration modules also exercise the module-level ASGI app and
    # leave its terminal runner cached after lifespan shutdown.
    main_module._knowledge_entity_processing_service = None
    try:
        with TestClient(main_module.app) as client:
            try:
                yield client
            finally:
                client.portal.call(main_module._unregister_service, main_module.app)
                client.portal.call(close_redis)
    finally:
        # ``stop()`` intentionally makes the application runner terminal. Each
        # test lifespan therefore needs a newly wired runner instead of reusing
        # the stopped module-level service from the previous TestClient.
        main_module._knowledge_entity_processing_service = None


def _assert_real_runtime_configuration() -> None:
    """Fail fast when the selected infrastructure runtime is incomplete."""

    settings = get_settings()
    required = {
        "DB_HOST": settings.db_host,
        "DB_USER": settings.db_user,
        "DB_PASS": settings.db_pass,
        "MINIO_ENDPOINT": settings.kb_minio_endpoint,
        "MINIO_ACCESS_KEY": settings.kb_minio_access_key,
        "MINIO_SECRET_KEY": settings.kb_minio_secret_key,
        "EMBEDDING_MODEL_NAME": settings.embedding_model_name,
        "EMBEDDING_DIMENSION": settings.embedding_dimension,
        "REDIS_HOST": settings.redis_host,
        "REDIS_PORT": settings.redis_port,
    }
    missing = [name for name, value in required.items() if value in (None, "", 0)]
    if missing:
        pytest.fail(
            "KnowledgeEntity integration requires configured infrastructure "
            f"dependencies; missing: {', '.join(missing)}",
            pytrace=False,
        )


def _assert_success(response) -> dict:
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["resultCode"] == "0", payload
    return payload["resultObject"]


def _assert_error(response, *, message: str | None = None) -> dict:
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["resultCode"] == "-1", payload
    if message is not None:
        assert message in payload["resultMsg"], payload
    return payload


def _create_kb(client: TestClient) -> str:
    result = _assert_success(
        client.post(
            "/api/v1/knowledgeBases/create",
            json={"knName": f"KnowledgeEntity real integration {uuid4().hex}"},
        )
    )
    return result["knCode"]


def _delete_kb(client: TestClient, kb_code: str) -> None:
    response = client.post(
        "/api/v1/knowledgeBases/delete",
        json={"knCode": kb_code},
    )
    _assert_success(response)


def _upload_markdown(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    content: str,
) -> None:
    _upload_file(
        client,
        kb_code=kb_code,
        file_path=file_path,
        content=content.encode("utf-8"),
        content_type="text/markdown",
    )


def _upload_file(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    content: bytes,
    content_type: str,
) -> None:
    _assert_success(
        client.post(
            "/api/v1/knowledgeItems/import",
            data={"knCode": kb_code, "filePath": file_path},
            files={
                "fileContent": (
                    file_path.rsplit("/", 1)[-1],
                    content,
                    content_type,
                )
            },
        )
    )


def _update_markdown(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    content: str,
) -> None:
    _assert_success(
        client.post(
            "/api/v1/knowledgeItems/update",
            data={"knCode": kb_code, "filePath": file_path},
            files={
                "fileContent": (
                    file_path.rsplit("/", 1)[-1],
                    content.encode("utf-8"),
                    "text/markdown",
                )
            },
        )
    )


def _build_markdown_index(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
) -> None:
    response = client.post(
        "/api/v1/fileToMarkdownIndex",
        json={"knCode": kb_code, "filePath": file_path},
    )
    _assert_success(response)
    build_status = _assert_success(
        client.post(
            "/api/v1/fileBuildStatus",
            json={"knCode": kb_code, "filePath": file_path},
        )
    )
    assert build_status["status"] == "complete", build_status
    build_result = _assert_success(
        client.post(
            "/api/v1/buildResult",
            json={
                "knCode": kb_code,
                "filePath": file_path,
                "includeMarkdown": True,
            },
        )
    )
    assert build_result["build"]["status"] == "complete"
    assert build_result["markdown"]["available"] is True
    assert build_result["markdown"]["data"]
    assert build_result["chunks"]["total"] >= 1
    assert build_result["embedding"]["embeddedChunkCount"] >= 1
    assert build_result["retrieval"]["indexedChunkCount"] >= 1


def _metadata(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    field_names: Iterable[str],
) -> dict:
    return _assert_success(
        client.post(
            "/api/v1/knowledgeItems/metadata/get",
            json={
                "knCode": kb_code,
                "filePath": file_path,
                "metadataFieldList": list(field_names),
            },
        )
    )["metadata"]


def _update_metadata(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    operations: list[dict[str, Any]],
) -> None:
    _assert_success(
        client.post(
            "/api/v1/knowledgeItems/metadata/update",
            json={
                "knCode": kb_code,
                "filePath": file_path,
                "operationList": operations,
            },
        )
    )


def _set_entity_metadata(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    entity_name: str,
    aliases: list[str] | None = None,
    document_kind: str = "knowledgeEntity",
) -> None:
    _update_metadata(
        client,
        kb_code=kb_code,
        file_path=file_path,
        operations=[
            {
                "propertyName": "documentKind",
                "operation": "set",
                "valueType": "string",
                "value": document_kind,
            },
            {
                "propertyName": "entityName",
                "operation": "set",
                "valueType": "string",
                "value": entity_name,
            },
            {
                "propertyName": "aliases",
                "operation": "set",
                "valueType": "stringList",
                "value": aliases or [],
            },
        ],
    )


def _eligibility(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    capability: str,
) -> dict:
    body: dict[str, Any] = {
        "knCode": kb_code,
        "filePath": file_path,
        "capability": capability,
    }
    return _assert_success(
        client.post(
            "/api/v1/knowledgeItems/processingEligibility",
            json=body,
        )
    )


def _read_markdown(client: TestClient, *, kb_code: str, file_path: str) -> str:
    return _assert_success(
        client.post(
            "/api/v1/readFile",
            json={"knCode": kb_code, "filePath": file_path},
        )
    )["data"]


def _download_file(client: TestClient, *, kb_code: str, file_path: str) -> bytes:
    response = client.post(
        "/api/v1/downloadFile",
        json={"knCode": kb_code, "filePath": file_path},
    )
    assert response.status_code == 200, response.text
    return response.content


def _semantic_relations(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    direction: str = "BOTH",
    relation_codes: list[str] | None = None,
) -> dict:
    body: dict[str, Any] = {
        "knCode": kb_code,
        "filePath": file_path,
        "direction": direction,
        "pageNum": 1,
        "pageSize": 500,
    }
    if relation_codes is not None:
        body["relationCodeList"] = relation_codes
    return _assert_success(
        client.post("/api/v1/knowledgeItems/semanticRelations", json=body)
    )


def _reference_result(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str,
    direction: str = "all",
) -> dict:
    return _assert_success(
        client.post(
            "/api/v1/knowledgeItems/references",
            json={
                "knCode": kb_code,
                "filePath": file_path,
                "direction": direction,
            },
        )
    )


def _move_file(
    client: TestClient,
    *,
    kb_code: str,
    source_path: str,
    target_file_path: str,
) -> None:
    result = _assert_success(
        client.post(
            "/api/v1/knowledgeItems/move",
            json={
                "knCode": kb_code,
                "sourcePath": [source_path],
                "targetFilePath": target_file_path,
            },
        )
    )
    assert result["data"] == [
        {
            "sourcePath": source_path,
            "targetPath": target_file_path,
            "success": True,
            "error": None,
        }
    ]


def _delete_file(client: TestClient, *, kb_code: str, file_path: str) -> None:
    _assert_success(
        client.post(
            "/api/v1/knowledgeItems/delete",
            json={"knCode": kb_code, "filePath": file_path},
        )
    )


def _task_page(
    client: TestClient,
    *,
    kb_code: str,
    file_path: str | None = None,
    batch_id: str | None = None,
    task_type: str | None = None,
    latest_only: bool = False,
    status_list: list[str] | None = None,
    include_details: bool = True,
    page_num: int = 1,
    page_size: int = 100,
) -> dict:
    body: dict[str, object] = {
        "knCode": kb_code,
        "latestOnly": latest_only,
        "includeDetails": include_details,
        "pageNum": page_num,
        "pageSize": page_size,
    }
    if file_path is not None:
        body["filePath"] = file_path
    if batch_id is not None:
        body["batchId"] = batch_id
    if task_type is not None:
        body["taskType"] = task_type
    if status_list is not None:
        body["statusList"] = status_list
    return _assert_success(
        client.post("/api/v1/knowledgeItems/processingTaskStatus", json=body)
    )


def _wait_for_batch(
    client: TestClient,
    *,
    kb_code: str,
    batch_id: str,
    expected_task_ids: Iterable[str],
    timeout_seconds: float = TASK_TIMEOUT_SECONDS,
) -> list[dict]:
    expected = set(expected_task_ids)
    assert expected, "the accepted batch must contain at least one real task"
    deadline = time.monotonic() + timeout_seconds
    latest: list[dict] = []
    while time.monotonic() < deadline:
        page = _task_page(
            client,
            kb_code=kb_code,
            batch_id=batch_id,
            latest_only=False,
        )
        latest = [item for item in page["data"] if item["taskId"] in expected]
        if len(latest) == len(expected) and all(
            item["status"] in TERMINAL_TASK_STATUSES for item in latest
        ):
            failures = [item for item in latest if item["status"] != "SUCCEEDED"]
            assert not failures, failures
            return latest
        time.sleep(2.0)
    pytest.fail(
        f"KnowledgeEntity batch {batch_id} did not finish within "
        f"{timeout_seconds}s; last task state: {latest}"
    )


def _entity_paths_from_tasks(tasks: Iterable[dict]) -> list[str]:
    paths: list[str] = []
    for task in tasks:
        for action in (task.get("result") or {}).get("actions", []):
            path = action.get("filePath")
            if action.get("action") in {"CREATED", "ANCHORED"} and path:
                paths.append(path)
    return list(dict.fromkeys(paths))


def _wait_for_batch_snapshots(
    client: TestClient,
    *,
    kb_code: str,
    batch_id: str,
    expected_task_ids: Iterable[str],
    require_success: bool = True,
    timeout_seconds: float = TASK_TIMEOUT_SECONDS,
) -> tuple[list[dict], list[list[dict]]]:
    expected = set(expected_task_ids)
    assert expected
    deadline = time.monotonic() + timeout_seconds
    snapshots: list[list[dict]] = []
    latest: list[dict] = []
    while time.monotonic() < deadline:
        page = _task_page(
            client,
            kb_code=kb_code,
            batch_id=batch_id,
            latest_only=False,
        )
        latest = [item for item in page["data"] if item["taskId"] in expected]
        if latest:
            snapshots.append(latest)
        if len(latest) == len(expected) and all(
            item["status"] in TERMINAL_TASK_STATUSES for item in latest
        ):
            if require_success:
                assert all(item["status"] == "SUCCEEDED" for item in latest), latest
            return latest, snapshots
        time.sleep(0.1)
    pytest.fail(
        f"KnowledgeEntity batch {batch_id} did not finish within "
        f"{timeout_seconds}s; last task state: {latest}"
    )


def _entity_markdown(
    *,
    entity_name: str,
    aliases: Iterable[str] = (),
    body: str | None = None,
) -> str:
    alias_values = ", ".join(f'"{value}"' for value in aliases)
    return f"""---
documentKind: knowledgeEntity
processingCapabilities: [entityEnrich]
entityName: {entity_name}
aliases: [{alias_values}]
---

# {entity_name}

{body or f"{entity_name} is a stable integration-test entity."}
"""


def _pdf_bytes(text: str) -> bytes:
    import fitz

    document = fitz.open()
    try:
        page = document.new_page()
        page.insert_text((72, 72), text)
        return document.tobytes()
    finally:
        document.close()


def _docx_bytes(text: str) -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph(text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


async def _db_fetch_all(
    query: str,
    params: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    connection = await build_connection_factory(get_settings())()
    try:
        cursor = connection.cursor()
        if params is None:
            await cursor.execute(query)
        else:
            await cursor.execute(query, params)
        return list(await cursor.fetchall())
    finally:
        await connection.close()


async def _db_execute(
    statements: str | Iterable[str],
    params: dict[str, Any] | None = None,
) -> None:
    connection = await build_connection_factory(get_settings())()
    try:
        cursor = connection.cursor()
        if isinstance(statements, str):
            statements = (statements,)
        for statement in statements:
            if params is None:
                await cursor.execute(statement)
            else:
                await cursor.execute(statement, params)
        await connection.commit()
    except Exception:
        await connection.rollback()
        raise
    finally:
        await connection.close()


def _db_rows(
    query: str,
    params: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    return asyncio.run(_db_fetch_all(query, params))


def _db_run(
    statements: str | Iterable[str],
    params: dict[str, Any] | None = None,
) -> None:
    asyncio.run(_db_execute(statements, params))


def _relation_assertion_rows(
    *, kb_code: str, source_path: str | None = None
) -> list[dict[str, Any]]:
    source_filter = ""
    params: dict[str, Any] = {"knowledge_base_id": int(kb_code)}
    if source_path is not None:
        source_filter = "AND source.virtual_path = %(source_path)s"
        params["source_path"] = source_path
    return _db_rows(
        f"""
        SELECT assertion.*, source.virtual_path AS source_path,
               target.virtual_path AS target_path
        FROM knowledge_file_reference assertion
        JOIN knowledge_fs_entry source
          ON source.kid = assertion.source_fs_entry_id
        LEFT JOIN knowledge_fs_entry target
          ON target.kid = assertion.target_fs_entry_id
        WHERE assertion.knowledge_base_id = %(knowledge_base_id)s
          {source_filter}
        ORDER BY assertion.kid
        """,
        params,
    )


def _latest_timeline(*, kb_code: str, file_path: str) -> dict[str, Any] | None:
    rows = _db_rows(
        """
        SELECT timeline.*
        FROM knowledge_file_update_timeline timeline
        JOIN knowledge_fs_entry entry ON entry.kid = timeline.fs_entry_id
        WHERE entry.knowledge_base_id = %(knowledge_base_id)s
          AND entry.virtual_path = %(file_path)s
        ORDER BY timeline.kid DESC
        LIMIT 1
        """,
        {"knowledge_base_id": int(kb_code), "file_path": file_path},
    )
    return rows[0] if rows else None


async def _repository_entity_surfaces(
    knowledge_base_id: int | None,
) -> list[dict[str, Any]]:
    connection = await build_connection_factory(get_settings())()
    try:
        return await KnowledgeEntityRepository().list_entity_surfaces(
            connection.cursor(), knowledge_base_id=knowledge_base_id
        )
    finally:
        await connection.close()


def _entity_surfaces(knowledge_base_id: int | None) -> list[dict[str, Any]]:
    return asyncio.run(_repository_entity_surfaces(knowledge_base_id))


@pytest.mark.integration
def test_knowledge_entity_real_api_end_to_end() -> None:
    """Exercise KnowledgeEntity through only real public APIs and dependencies."""

    _assert_real_runtime_configuration()
    source_path = "/documents/zetaknowledge-protocol.md"
    second_source_path = "/documents/atlas-index.md"
    source_v1 = """# ZetaKnowledge Protocol

ZetaKnowledge Protocol is an internal knowledge-governance protocol. It defines
how AtlasIndex records canonical knowledge entities and their aliases.

AtlasIndex is the protocol's canonical entity index. ZetaKnowledge Protocol
depends on AtlasIndex for identity resolution.
"""
    source_v2 = (
        source_v1 + "\nThe protocol is maintained by the Knowledge Platform team.\n"
    )
    second_source = """# AtlasIndex

AtlasIndex is the canonical entity index used by ZetaKnowledge Protocol. It
stores stable entity names and aliases for knowledge-governance workflows.
"""

    with _real_test_client() as client:
        kb_code = _create_kb(client)
        try:
            # Import and update must materialize/preserve the default original
            # document kind without front matter supplied by the caller.
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=source_path,
                content=source_v1,
            )
            imported_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=source_path,
                field_names=["documentKind", "processingCapabilities"],
            )
            assert imported_metadata["documentKind"]["value"] == "original"

            _update_markdown(
                client,
                kb_code=kb_code,
                file_path=source_path,
                content=source_v2,
            )
            updated_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=source_path,
                field_names=["documentKind"],
            )
            assert updated_metadata["documentKind"]["value"] == "original"

            _build_markdown_index(client, kb_code=kb_code, file_path=source_path)
            eligibility = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/processingEligibility",
                    json={
                        "knCode": kb_code,
                        "filePath": source_path,
                        "capability": "entityDiscovery",
                    },
                )
            )
            assert eligibility["documentKind"] == "original"
            assert eligibility["eligibility"] == "ELIGIBLE_AND_STALE"
            assert eligibility["reasonCode"] == "NEVER_PROCESSED"

            single_batch = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityDiscovery",
                    json={
                        "knCode": kb_code,
                        "filePath": source_path,
                        "maxEntities": 3,
                    },
                )
            )
            assert single_batch["scope"] == "SINGLE_FILE"
            assert single_batch["acceptedCount"] == 1
            single_tasks = _wait_for_batch(
                client,
                kb_code=kb_code,
                batch_id=single_batch["batchId"],
                expected_task_ids=(item["taskId"] for item in single_batch["tasks"]),
            )

            # Querying by KB id + file path must return only that source file and
            # include the durable result written by the asynchronous worker.
            source_status = _task_page(
                client,
                kb_code=kb_code,
                file_path=source_path,
                task_type="ENTITY_DISCOVERY",
                latest_only=True,
            )
            assert source_status["filePath"] == source_path
            assert source_status["total"] == 1
            assert source_status["data"][0]["filePath"] == source_path
            assert source_status["data"][0]["status"] == "SUCCEEDED"
            assert source_status["data"][0]["result"]

            entity_paths = _entity_paths_from_tasks(single_tasks)
            assert entity_paths, single_tasks
            assert all(path.startswith("/KnowledgeEntity/") for path in entity_paths)
            entity_path = entity_paths[0]
            assert not re.search(r"-[0-9a-f]{12}\.md$", Path(entity_path).name)
            entity_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                field_names=[
                    "documentKind",
                    "processingCapabilities",
                    "entityName",
                    "aliases",
                ],
            )
            assert entity_metadata["documentKind"]["value"] == "knowledgeEntity"
            assert "entityEnrich" in entity_metadata["processingCapabilities"]["value"]
            assert entity_metadata["entityName"]["value"]
            entity_markdown_before_enrich = _read_markdown(
                client, kb_code=kb_code, file_path=entity_path
            )
            assert entity_markdown_before_enrich.strip()
            assert source_path in entity_markdown_before_enrich
            entity_directory = _assert_success(
                client.post(
                    "/api/v1/listDir",
                    json={
                        "knCode": kb_code,
                        "directoryPath": "/KnowledgeEntity",
                    },
                )
            )
            assert any(
                item.get("filePath") == entity_path
                or item.get("path") == entity_path
                or item.get("name") == entity_path
                for item in entity_directory["data"]
            )

            mention_page = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/semanticRelations",
                    json={
                        "knCode": kb_code,
                        "filePath": source_path,
                        "direction": "OUTGOING",
                        "relationCodeList": ["MENTIONS"],
                    },
                )
            )
            assert mention_page["total"] >= 1
            assert any(
                item["target"]["filePath"] == entity_path
                and item["relationCode"] == "MENTIONS"
                for item in mention_page["data"]
            )
            entity_outgoing_mentions = _semantic_relations(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                direction="OUTGOING",
                relation_codes=["MENTIONS"],
            )
            assert any(
                item["target"]["filePath"] == source_path
                and item["relationCode"] == "MENTIONS"
                for item in entity_outgoing_mentions["data"]
            )
            entity_assertions = _relation_assertion_rows(
                kb_code=kb_code,
                source_path=entity_path,
            )
            assert any(
                item["relation_code"] == "MENTIONS"
                and item["discovered_by"] == "MARKDOWN_PARSER"
                and item["target_path"] == source_path
                for item in entity_assertions
            )
            assert not any(
                item["discovered_by"] == "ENTITY_DISCOVERY"
                for item in entity_assertions
            )

            # An unbuilt second original document and the fresh first document
            # exercise all-KB enumeration, eligibility, exclusion of generated
            # KnowledgeEntity files, and durable task reuse without duplicating
            # an already proven deterministic discovery-model call.
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=second_source_path,
                content=second_source,
            )
            whole_batch = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityDiscovery",
                    json={"knCode": kb_code, "maxEntities": 2},
                )
            )
            assert whole_batch["scope"] == "WHOLE_KB"
            assert whole_batch["eligibleCount"] == 1
            assert whole_batch["acceptedCount"] == 0
            assert whole_batch["reusedCount"] == 1
            assert whole_batch["skippedCount"] >= 1
            reused_tasks = [
                item for item in whole_batch["tasks"] if item["reused"] is True
            ]
            assert len(reused_tasks) == 1
            assert reused_tasks[0]["filePath"] == source_path
            assert {item["filePath"] for item in whole_batch["tasks"]} >= {
                source_path,
                second_source_path,
                entity_path,
            }

            discovery_fresh = _eligibility(
                client,
                kb_code=kb_code,
                file_path=source_path,
                capability="entityDiscovery",
            )
            assert discovery_fresh["eligibility"] == "ELIGIBLE_BUT_FRESH"
            assert discovery_fresh["reasonCode"] == "INPUT_UNCHANGED"

            source_v3 = source_v2 + "\nAtlasIndex remains the canonical registry.\n"
            _update_markdown(
                client,
                kb_code=kb_code,
                file_path=source_path,
                content=source_v3,
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=source_path)
            discovery_stale = _eligibility(
                client,
                kb_code=kb_code,
                file_path=source_path,
                capability="entityDiscovery",
            )
            assert discovery_stale["eligibility"] == "ELIGIBLE_AND_STALE"
            assert discovery_stale["reasonCode"] == "INPUT_CHANGED"
            forced_discovery = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityDiscovery",
                    json={
                        "knCode": kb_code,
                        "filePath": source_path,
                        "maxEntities": 3,
                        "force": True,
                    },
                )
            )
            assert forced_discovery["acceptedCount"] == 1
            assert forced_discovery["reusedCount"] == 0
            assert forced_discovery["tasks"][0]["taskId"] != single_tasks[0]["taskId"]
            _wait_for_batch(
                client,
                kb_code=kb_code,
                batch_id=forced_discovery["batchId"],
                expected_task_ids=(
                    item["taskId"] for item in forced_discovery["tasks"]
                ),
            )
            enrich_eligibility = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/processingEligibility",
                    json={
                        "knCode": kb_code,
                        "filePath": entity_path,
                        "capability": "entityEnrich",
                    },
                )
            )
            assert enrich_eligibility["documentKind"] == "knowledgeEntity"
            assert enrich_eligibility["eligibility"] == "ELIGIBLE_AND_STALE"
            before_enrich_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                field_names=["fileSignature", "entityName"],
            )
            assert _latest_timeline(kb_code=kb_code, file_path=entity_path) is None

            enrich_batch = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityEnrich",
                    json={
                        "knCode": kb_code,
                        "filePath": entity_path,
                        "topK": 5,
                    },
                )
            )
            assert enrich_batch["scope"] == "SINGLE_FILE"
            assert enrich_batch["acceptedCount"] == 1
            enrich_tasks = _wait_for_batch(
                client,
                kb_code=kb_code,
                batch_id=enrich_batch["batchId"],
                expected_task_ids=(item["taskId"] for item in enrich_batch["tasks"]),
            )
            assert enrich_tasks[0]["taskType"] == "DOCUMENT_ENRICH"
            assert enrich_tasks[0]["result"]["actions"][0]["action"] == "UPDATED"
            assert "templateCoverage" in enrich_tasks[0]["result"]
            assert "missingSections" in enrich_tasks[0]["result"]
            assert "warnings" in enrich_tasks[0]["result"]
            enriched_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                field_names=["documentKind", "entityName"],
            )
            assert enriched_metadata["documentKind"]["value"] == "knowledgeEntity"
            assert enriched_metadata["entityName"]["value"]
            after_enrich_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                field_names=["fileSignature", "entityName"],
            )
            assert (
                after_enrich_metadata["fileSignature"]["value"]
                != before_enrich_metadata["fileSignature"]["value"]
            )
            assert (
                after_enrich_metadata["entityName"]
                == before_enrich_metadata["entityName"]
            )
            enriched_markdown = _read_markdown(
                client, kb_code=kb_code, file_path=entity_path
            )
            assert enriched_markdown.strip()
            assert enriched_metadata["entityName"]["value"] in enriched_markdown
            timeline = _latest_timeline(kb_code=kb_code, file_path=entity_path)
            assert timeline is not None
            assert timeline["event_type"] == "UPDATE"
            assert (
                timeline["old_checksum"]
                == before_enrich_metadata["fileSignature"]["value"]
            )
            assert (
                timeline["new_checksum"]
                == after_enrich_metadata["fileSignature"]["value"]
            )

            enrich_fresh = _eligibility(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                capability="entityEnrich",
            )
            assert enrich_fresh["eligibility"] == "ELIGIBLE_BUT_FRESH"
            assert enrich_fresh["reasonCode"] == "NO_NEW_RELATIONS"

            # Enrichment replaces only this entity's generated outgoing edges;
            # the endpoint must remain queryable even when the model emits none.
            enriched_relations = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/semanticRelations",
                    json={
                        "knCode": kb_code,
                        "filePath": entity_path,
                        "direction": "BOTH",
                    },
                )
            )
            assert enriched_relations["fileId"] == enrich_eligibility["fileId"]
            assert enriched_relations["total"] >= 1
        finally:
            _delete_kb(client, kb_code)


@pytest.mark.integration
def test_knowledge_entity_real_eligibility_and_request_matrix() -> None:
    """Cover KE-M2/M3/M5-M9, KE-D8, KE-E4/E8, and KE-T3/T6."""

    _assert_real_runtime_configuration()
    with _real_test_client() as client:
        kb_code = _create_kb(client)
        try:
            # KE-T3: a KB and a live file with no semantic task both return an
            # empty page, while a missing path is a documented validation error.
            assert _task_page(client, kb_code=kb_code)["data"] == []
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path="/untasked.md",
                content="# Untasked\n",
            )
            untasked_page = _task_page(
                client, kb_code=kb_code, file_path="/untasked.md"
            )
            assert untasked_page["total"] == 0
            assert untasked_page["data"] == []
            _assert_error(
                client.post(
                    "/api/v1/knowledgeItems/processingTaskStatus",
                    json={"knCode": kb_code, "filePath": "/missing.md"},
                ),
                message="document not found",
            )

            # KE-M2/M9: direct imports below the reserved directory receive the
            # entity kind immediately. Readiness precedes identity validation.
            direct_entity_path = "/KnowledgeEntity/direct-import.md"
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=direct_entity_path,
                content="# Direct import\n",
            )
            direct_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=direct_entity_path,
                field_names=["documentKind", "processingCapabilities"],
            )
            assert direct_metadata == {
                "documentKind": {"valueType": "string", "value": "knowledgeEntity"}
            }
            before_build = _eligibility(
                client,
                kb_code=kb_code,
                file_path=direct_entity_path,
                capability="entityEnrich",
            )
            assert before_build["reasonCode"] == "CONTENT_NOT_READY"
            _build_markdown_index(client, kb_code=kb_code, file_path=direct_entity_path)
            identity_incomplete = _eligibility(
                client,
                kb_code=kb_code,
                file_path=direct_entity_path,
                capability="entityEnrich",
            )
            assert identity_incomplete["reasonCode"] == "IDENTITY_METADATA_INCOMPLETE"

            # KE-M3/M8: an explicit kind survives document update, and Enrich
            # rejects an otherwise complete entity outside the reserved path.
            outside_entity_path = "/outside/explicit-entity.md"
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=outside_entity_path,
                content="# Explicit entity\n",
            )
            _set_entity_metadata(
                client,
                kb_code=kb_code,
                file_path=outside_entity_path,
                entity_name="ExplicitBoundaryEntity",
                aliases=["BoundaryAlias"],
            )
            _update_markdown(
                client,
                kb_code=kb_code,
                file_path=outside_entity_path,
                content="# Explicit entity updated\n",
            )
            explicit_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=outside_entity_path,
                field_names=["documentKind", "entityName", "aliases"],
            )
            assert explicit_metadata["documentKind"]["value"] == "knowledgeEntity"
            assert explicit_metadata["entityName"]["value"] == "ExplicitBoundaryEntity"
            _build_markdown_index(
                client, kb_code=kb_code, file_path=outside_entity_path
            )
            outside_eligibility = _eligibility(
                client,
                kb_code=kb_code,
                file_path=outside_entity_path,
                capability="entityEnrich",
            )
            assert outside_eligibility["reasonCode"] == (
                "KNOWLEDGE_ENTITY_PATH_REQUIRED"
            )

            # KE-M5: missing processingCapabilities applies the kind default;
            # an explicit empty list disables it; unset restores the default.
            capability_path = "/matrix/capability.md"
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=capability_path,
                content="# Capability matrix\nStable text content.\n",
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=capability_path)
            default_capability = _eligibility(
                client,
                kb_code=kb_code,
                file_path=capability_path,
                capability="entityDiscovery",
            )
            assert default_capability["reasonCode"] == "NEVER_PROCESSED"
            _update_metadata(
                client,
                kb_code=kb_code,
                file_path=capability_path,
                operations=[
                    {
                        "propertyName": "processingCapabilities",
                        "operation": "set",
                        "valueType": "stringList",
                        "value": [],
                    }
                ],
            )
            disabled_capability = _eligibility(
                client,
                kb_code=kb_code,
                file_path=capability_path,
                capability="entityDiscovery",
            )
            assert disabled_capability["reasonCode"] == "CAPABILITY_DISABLED"
            _update_metadata(
                client,
                kb_code=kb_code,
                file_path=capability_path,
                operations=[
                    {
                        "propertyName": "processingCapabilities",
                        "operation": "unset",
                    }
                ],
            )
            restored_capability = _eligibility(
                client,
                kb_code=kb_code,
                file_path=capability_path,
                capability="entityDiscovery",
            )
            assert restored_capability["reasonCode"] == "NEVER_PROCESSED"

            # KE-M6: every text suffix crosses the format gate after a real
            # build, while PDF/Office remain ineligible despite valid sidecars.
            text_documents = {
                "/formats/a.md": (b"# Markdown\n", "text/markdown"),
                "/formats/b.markdown": (b"# Markdown long\n", "text/markdown"),
                "/formats/c.txt": (b"plain text\n", "text/plain"),
                "/formats/d.html": (
                    b"<html><body><h1>HTML</h1></body></html>",
                    "text/html",
                ),
                "/formats/e.htm": (
                    b"<html><body>HTM content</body></html>",
                    "text/html",
                ),
                "/formats/f.csv": (b"name,value\nalpha,1\n", "text/csv"),
            }
            for file_path, (content, content_type) in text_documents.items():
                _upload_file(
                    client,
                    kb_code=kb_code,
                    file_path=file_path,
                    content=content,
                    content_type=content_type,
                )
                _build_markdown_index(client, kb_code=kb_code, file_path=file_path)
                evaluation = _eligibility(
                    client,
                    kb_code=kb_code,
                    file_path=file_path,
                    capability="entityDiscovery",
                )
                assert evaluation["reasonCode"] == "NEVER_PROCESSED", evaluation

            binary_documents = {
                "/formats/g.pdf": (
                    _pdf_bytes("Binary PDF sidecar eligibility"),
                    "application/pdf",
                ),
                "/formats/h.docx": (
                    _docx_bytes("Binary Office sidecar eligibility"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            }
            for file_path, (content, content_type) in binary_documents.items():
                _upload_file(
                    client,
                    kb_code=kb_code,
                    file_path=file_path,
                    content=content,
                    content_type=content_type,
                )
                _build_markdown_index(client, kb_code=kb_code, file_path=file_path)
                evaluation = _eligibility(
                    client,
                    kb_code=kb_code,
                    file_path=file_path,
                    capability="entityDiscovery",
                )
                assert evaluation["reasonCode"] == "UNSUPPORTED_FILE_FORMAT"

            # KE-M7: non-text MIME and suffix precedence are implemented here;
            # the text/* fallback contract is isolated in a strict xfail below.
            _upload_file(
                client,
                kb_code=kb_code,
                file_path="/formats/no-suffix-binary",
                content=b"binary",
                content_type="application/octet-stream",
            )
            suffixless_binary = _eligibility(
                client,
                kb_code=kb_code,
                file_path="/formats/no-suffix-binary",
                capability="entityDiscovery",
            )
            assert suffixless_binary["reasonCode"] == "UNSUPPORTED_FILE_FORMAT"
            _upload_file(
                client,
                kb_code=kb_code,
                file_path="/formats/text-disguised.pdf",
                content=b"not a pdf",
                content_type="text/plain",
            )
            disguised_pdf = _eligibility(
                client,
                kb_code=kb_code,
                file_path="/formats/text-disguised.pdf",
                capability="entityDiscovery",
            )
            assert disguised_pdf["reasonCode"] == "UNSUPPORTED_FILE_FORMAT"

            # KE-M8/M9/E4: Enrich rejects non-Markdown content before identity,
            # then rejects a complete Markdown identity that has no evidence.
            for file_path, content, content_type in (
                (
                    "/KnowledgeEntity/not-markdown.txt",
                    b"entity text",
                    "text/plain",
                ),
                (
                    "/KnowledgeEntity/not-markdown.pdf",
                    _pdf_bytes("entity pdf"),
                    "application/pdf",
                ),
            ):
                _upload_file(
                    client,
                    kb_code=kb_code,
                    file_path=file_path,
                    content=content,
                    content_type=content_type,
                )
                _build_markdown_index(client, kb_code=kb_code, file_path=file_path)
                evaluation = _eligibility(
                    client,
                    kb_code=kb_code,
                    file_path=file_path,
                    capability="entityEnrich",
                )
                assert evaluation["reasonCode"] == "UNSUPPORTED_CONTENT_TYPE"

            no_evidence_path = "/KnowledgeEntity/no-evidence.md"
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=no_evidence_path,
                content=_entity_markdown(entity_name="NoEvidenceEntity"),
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=no_evidence_path)
            no_evidence_before = _metadata(
                client,
                kb_code=kb_code,
                file_path=no_evidence_path,
                field_names=["fileSignature", "entityName"],
            )
            no_evidence_markdown = _read_markdown(
                client, kb_code=kb_code, file_path=no_evidence_path
            )
            no_evidence = _eligibility(
                client,
                kb_code=kb_code,
                file_path=no_evidence_path,
                capability="entityEnrich",
            )
            assert no_evidence["reasonCode"] == "NO_EVIDENCE"
            skipped_enrich = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityEnrich",
                    json={"knCode": kb_code, "filePath": no_evidence_path},
                )
            )
            assert skipped_enrich["acceptedCount"] == 0
            assert skipped_enrich["skippedCount"] == 1
            assert (
                _metadata(
                    client,
                    kb_code=kb_code,
                    file_path=no_evidence_path,
                    field_names=[
                        "fileSignature",
                        "entityName",
                    ],
                )
                == no_evidence_before
            )
            assert (
                _read_markdown(client, kb_code=kb_code, file_path=no_evidence_path)
                == no_evidence_markdown
            )
            assert _latest_timeline(kb_code=kb_code, file_path=no_evidence_path) is None
            assert not _relation_assertion_rows(
                kb_code=kb_code, source_path=no_evidence_path
            )

            # KE-D8/E8/T6: strict HTTP models reject every removed target,
            # template, and callback field without creating a task.
            task_count_before_invalid = _task_page(client, kb_code=kb_code)["total"]
            invalid_requests = (
                (
                    "/api/v1/knowledgeItems/entityDiscovery",
                    {
                        "knCode": kb_code,
                        "filePath": capability_path,
                        "targetKnCode": "other",
                    },
                ),
                (
                    "/api/v1/knowledgeItems/entityDiscovery",
                    {
                        "knCode": kb_code,
                        "filePath": capability_path,
                        "targetDirectoryPath": "/custom",
                    },
                ),
                (
                    "/api/v1/knowledgeItems/entityDiscovery",
                    {
                        "knCode": kb_code,
                        "filePath": capability_path,
                        "callback": "https://example.invalid/callback",
                    },
                ),
                (
                    "/api/v1/knowledgeItems/entityEnrich",
                    {
                        "knCode": kb_code,
                        "filePath": no_evidence_path,
                        "template": "# client template",
                    },
                ),
                (
                    "/api/v1/knowledgeItems/entityEnrich",
                    {
                        "knCode": kb_code,
                        "filePath": no_evidence_path,
                        "targetKnCode": "other",
                    },
                ),
                (
                    "/api/v1/knowledgeItems/entityEnrich",
                    {
                        "knCode": kb_code,
                        "filePath": no_evidence_path,
                        "targetDirectoryPath": "/custom",
                    },
                ),
                (
                    "/api/v1/knowledgeItems/entityEnrich",
                    {
                        "knCode": kb_code,
                        "filePath": no_evidence_path,
                        "callback": {"type": "callable"},
                    },
                ),
                (
                    "/api/v1/knowledgeItems/entityEnrich",
                    {
                        "knCode": kb_code,
                        "filePath": no_evidence_path,
                        "evidenceKnCodeList": [""],
                    },
                ),
                (
                    "/api/v1/knowledgeItems/entityEnrich",
                    {
                        "knCode": kb_code,
                        "filePath": no_evidence_path,
                        "evidenceKnCodeList": [kb_code, kb_code],
                    },
                ),
            )
            for route, body in invalid_requests:
                _assert_error(
                    client.post(route, json=body), message="request validation failed"
                )
            assert (
                _task_page(client, kb_code=kb_code)["total"]
                == task_count_before_invalid
            )
        finally:
            _delete_kb(client, kb_code)


@pytest.mark.integration
def test_knowledge_entity_real_anchor_relations_tasks_and_callbacks() -> None:
    """Cover real anchoring, task reuse/callbacks, and unified relations."""

    _assert_real_runtime_configuration()
    token = uuid4().hex[:10]
    entity_name = f"AnchorEntity{token}"
    entity_alias = f"AnchorAlias{token}"
    entity_path = f"/KnowledgeEntity/{entity_name}.md"
    foreign_entity_path = entity_path
    source_path = f"/sources/anchor-{token}.md"
    second_source_path = f"/sources/anchor-peer-{token}.md"
    with _real_test_client() as client:
        kb_code = _create_kb(client)
        foreign_kb_code = _create_kb(client)
        try:
            # KE-ENV3/D5: the repository executes both optional-scope branches
            # against real OpenGauss. The same system-wide surface exists in a
            # second KB, but only the subject KB entity may be anchored.
            for current_kb, current_path in (
                (kb_code, entity_path),
                (foreign_kb_code, foreign_entity_path),
            ):
                _upload_markdown(
                    client,
                    kb_code=current_kb,
                    file_path=current_path,
                    content=_entity_markdown(
                        entity_name=entity_name,
                        aliases=[entity_alias],
                    ),
                )
                _build_markdown_index(
                    client, kb_code=current_kb, file_path=current_path
                )

            global_surfaces = _entity_surfaces(None)
            local_surfaces = _entity_surfaces(int(kb_code))
            matching_global = [
                item for item in global_surfaces if item["entity_name"] == entity_name
            ]
            matching_local = [
                item for item in local_surfaces if item["entity_name"] == entity_name
            ]
            assert len(matching_global) == 2
            assert len(matching_local) == 1
            local_entity_id = str(matching_local[0]["kid"])
            assert matching_local[0]["file_path"] == entity_path

            source_markdown = f"""# Anchoring {token}

{entity_name} is also called {entity_alias}. The source depends on this stable
entity and links to its canonical file.

[{entity_name} details]({entity_path})
"""
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=source_path,
                content=source_markdown,
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=source_path)
            markdown_assertions = _relation_assertion_rows(
                kb_code=kb_code, source_path=source_path
            )
            assert len(markdown_assertions) == 1
            assert markdown_assertions[0]["discovered_by"] == "MARKDOWN_PARSER"
            assert str(markdown_assertions[0]["target_fs_entry_id"]) == local_entity_id

            # KE-D6/D7/ENV4: concurrent real HTTP requests serialize on the
            # source row. One creates a PENDING task and the other records a
            # skipped response that points at the active work through reuse.
            def request_discovery() -> dict:
                return _assert_success(
                    client.post(
                        "/api/v1/knowledgeItems/entityDiscovery",
                        json={
                            "knCode": kb_code,
                            "filePath": source_path,
                            "maxEntities": 1,
                        },
                    )
                )

            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                responses = list(executor.map(lambda _: request_discovery(), range(2)))
            assert sorted(item["acceptedCount"] for item in responses) == [0, 1]
            assert sorted(item["reusedCount"] for item in responses) == [0, 1]
            task_ids = {
                task["taskId"] for response in responses for task in response["tasks"]
            }
            assert len(task_ids) == 2
            accepted = next(item for item in responses if item["acceptedCount"] == 1)
            reused = next(item for item in responses if item["reusedCount"] == 1)
            assert accepted["tasks"][0]["status"] == "PENDING"
            assert reused["tasks"][0]["status"] == "SKIPPED"
            assert reused["tasks"][0]["reused"] is True
            accepted_task_ids = {item["taskId"] for item in accepted["tasks"]}
            first_tasks, snapshots = _wait_for_batch_snapshots(
                client,
                kb_code=kb_code,
                batch_id=accepted["batchId"],
                expected_task_ids=accepted_task_ids,
            )
            assert first_tasks[0]["status"] == "SUCCEEDED"
            assert first_tasks[0]["currentStage"] == "completed"
            assert first_tasks[0]["progress"] == 100
            assert first_tasks[0]["startedAt"]
            assert first_tasks[0]["finishedAt"]
            assert snapshots

            actions = first_tasks[0]["result"]["actions"]
            assert any(
                action["action"] in {"CREATED", "ANCHORED"}
                and str(action["entityFileId"]) == local_entity_id
                for action in actions
            )
            entity_count = _db_rows(
                """
                SELECT COUNT(*) AS total
                FROM knowledge_fs_entry entry
                JOIN knowledge_file_metadata_value metadata
                  ON metadata.fs_entry_id = entry.kid
                 AND metadata.property_name = 'entityName'
                 AND metadata.is_deleted = FALSE
                WHERE entry.knowledge_base_id = %(knowledge_base_id)s
                  AND entry.is_deleted = FALSE
                  AND metadata.value_string = %(entity_name)s
                """,
                {
                    "knowledge_base_id": int(kb_code),
                    "entity_name": entity_name,
                },
            )[0]["total"]
            assert int(entity_count) == 1

            # KE-R1/R10: a Markdown assertion and a discovery assertion fold
            # into one logical MENTIONS edge with two lightweight producers.
            first_relation_page = _semantic_relations(
                client,
                kb_code=kb_code,
                file_path=source_path,
                direction="OUTGOING",
                relation_codes=["MENTIONS"],
            )
            matching_relations = [
                item
                for item in first_relation_page["data"]
                if item["target"]["filePath"] == entity_path
            ]
            assert len(matching_relations) == 1
            first_relation = matching_relations[0]
            assert first_relation["assertionCount"] == 2
            assert first_relation["source"]["filePath"] == source_path
            relation_id = first_relation["relationId"]
            incoming = _semantic_relations(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                direction="INCOMING",
                relation_codes=["MENTIONS"],
            )
            assert any(
                item["relationId"] == relation_id
                and item["source"]["filePath"] == source_path
                for item in incoming["data"]
            )
            assert not any(
                item["source"]["fileId"] == item["target"]["fileId"]
                for item in first_relation_page["data"] + incoming["data"]
            )
            references = _reference_result(
                client,
                kb_code=kb_code,
                file_path=source_path,
                direction="outbound",
            )
            assert any(
                item.get("targetPath") == entity_path
                or item.get("filePath") == entity_path
                for item in references["outbound"]
            )

            # KE-T4: use the same application-scoped real service as an SDK
            # caller. The configured Protocol receives only committed file and
            # batch terminal events; status is then verified through HTTP.
            published_events: list[Any] = []

            class CapturingPublisher:
                async def publish(self, event: Any) -> None:
                    published_events.append(event)

            async def sdk_discovery(event_publisher) -> Any:
                service = (
                    await main_module.resolve_knowledge_entity_processing_service()
                )
                service.event_publisher_invoker.publisher = event_publisher
                return await service.discover_knowledge_entities(
                    EntityDiscoveryRequest(
                        knCode=kb_code,
                        filePath=source_path,
                        maxEntities=1,
                        force=True,
                    ),
                )

            callback_batch = client.portal.call(sdk_discovery, CapturingPublisher())
            assert callback_batch.accepted_count == 1
            callback_task_id = callback_batch.tasks[0].task_id
            callback_tasks = _wait_for_batch(
                client,
                kb_code=kb_code,
                batch_id=callback_batch.batch_id,
                expected_task_ids=[callback_task_id],
            )
            assert callback_tasks[0]["status"] == "SUCCEEDED"
            callback_events = [
                event
                for event in published_events
                if event.event_type == "semantic.discovery.file.completed"
                and event.payload.task_id == callback_task_id
            ]
            assert len(callback_events) == 1
            assert callback_events[0].payload.status == "SUCCEEDED"
            assert callback_events[0].payload.progress.completed_count == 1
            assert all(
                event.payload.batch_id == callback_batch.batch_id
                for event in callback_events
            )
            assert any(
                event.event_type == "semantic.discovery.batch.completed"
                and event.payload.batch_id == callback_batch.batch_id
                for event in published_events
            )

            # KE-R2: force creates a new producer run but replaces the old
            # discovery assertion. The logical relation remains stable.
            after_force_assertions = _relation_assertion_rows(
                kb_code=kb_code, source_path=source_path
            )
            assert sorted(row["discovered_by"] for row in after_force_assertions) == [
                "ENTITY_DISCOVERY",
                "MARKDOWN_PARSER",
            ]
            after_force_page = _semantic_relations(
                client,
                kb_code=kb_code,
                file_path=source_path,
                direction="OUTGOING",
                relation_codes=["MENTIONS"],
            )
            after_force_relation = next(
                item
                for item in after_force_page["data"]
                if item["target"]["filePath"] == entity_path
            )
            assert after_force_relation["relationId"] == relation_id
            assert after_force_relation["assertionCount"] == 2

            # KE-T5: callback failure is isolated from the task transaction.
            class RaisingPublisher:
                async def publish(self, event: Any) -> None:
                    del event
                    raise RuntimeError("intentional integration callback failure")

            failing_callback_batch = client.portal.call(
                sdk_discovery, RaisingPublisher()
            )
            failing_callback_tasks = _wait_for_batch(
                client,
                kb_code=kb_code,
                batch_id=failing_callback_batch.batch_id,
                expected_task_ids=[failing_callback_batch.tasks[0].task_id],
            )
            assert failing_callback_tasks[0]["status"] == "SUCCEEDED"
            assert _read_markdown(
                client, kb_code=kb_code, file_path=entity_path
            ).strip()

            # KE-T1/T2: real SQL combination filters and stable pagination.
            all_tasks = _task_page(
                client,
                kb_code=kb_code,
                task_type="ENTITY_DISCOVERY",
                latest_only=False,
            )
            source_tasks = _task_page(
                client,
                kb_code=kb_code,
                file_path=source_path,
                task_type="ENTITY_DISCOVERY",
                latest_only=False,
            )
            assert all_tasks["total"] == source_tasks["total"] == 4
            page_one = _task_page(
                client,
                kb_code=kb_code,
                file_path=source_path,
                task_type="ENTITY_DISCOVERY",
                status_list=["SUCCEEDED"],
                latest_only=False,
                include_details=False,
                page_num=1,
                page_size=1,
            )
            page_two = _task_page(
                client,
                kb_code=kb_code,
                file_path=source_path,
                task_type="ENTITY_DISCOVERY",
                status_list=["SUCCEEDED"],
                latest_only=False,
                include_details=False,
                page_num=2,
                page_size=1,
            )
            assert page_one["total"] == page_two["total"] == 3
            assert page_one["data"][0]["taskId"] != page_two["data"][0]["taskId"]
            assert "result" not in page_one["data"][0]
            assert "error" not in page_one["data"][0]
            batch_filtered = _task_page(
                client,
                kb_code=kb_code,
                batch_id=callback_batch.batch_id,
                task_type="ENTITY_DISCOVERY",
                status_list=["SUCCEEDED"],
                latest_only=False,
            )
            assert batch_filtered["total"] == 1
            assert batch_filtered["data"][0]["taskId"] == callback_task_id
            persisted_requests = _db_rows(
                """
                SELECT request_params
                FROM knowledge_semantic_processing_task
                WHERE knowledge_base_id = %(knowledge_base_id)s
                ORDER BY kid
                """,
                {"knowledge_base_id": int(kb_code)},
            )
            assert persisted_requests
            assert all(
                "callback" not in (row["request_params"] or {})
                for row in persisted_requests
            )

            # KE-R11: assert the deployed schema keeps only lightweight
            # assertion evidence and has no legacy heavy evidence table.
            columns = {
                row["column_name"]
                for row in _db_rows(
                    """
                    SELECT column_name
                    FROM information_schema.columns
                    WHERE table_schema = current_schema()
                      AND table_name = 'knowledge_file_reference'
                    """
                )
            }
            assert {
                "discovered_by",
                "producer_run_id",
                "evidence_fingerprint",
                "start_line",
                "end_line",
                "start_offset",
                "end_offset",
                "source_task_id",
            } <= columns
            assert not any("evidence_json" in column for column in columns)
            heavy_table_count = _db_rows(
                """
                SELECT COUNT(*) AS total
                FROM information_schema.tables
                WHERE table_schema = current_schema()
                  AND table_name = 'knowledge_document_relation_evidence'
                """
            )[0]["total"]
            assert int(heavy_table_count) == 0

            # KE-R5-R7: a second source owns an independent assertion to the
            # same target. Moving either endpoint preserves stable file IDs and
            # the logical relation ID; updating A rebuilds only A's outgoing
            # assertions and leaves B untouched.
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=second_source_path,
                content=(
                    f"# Peer source\n\n[{entity_name}]({entity_path}) is shared.\n"
                ),
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=second_source_path)
            peer_before = _semantic_relations(
                client,
                kb_code=kb_code,
                file_path=second_source_path,
                direction="OUTGOING",
                relation_codes=["MENTIONS"],
            )["data"][0]
            peer_relation_id = peer_before["relationId"]

            moved_entity_path = f"/KnowledgeEntity/moved-anchor-{token}.md"
            _move_file(
                client,
                kb_code=kb_code,
                source_path=entity_path,
                target_file_path=moved_entity_path,
            )
            moved_target_relation = next(
                item
                for item in _semantic_relations(
                    client,
                    kb_code=kb_code,
                    file_path=source_path,
                    direction="OUTGOING",
                    relation_codes=["MENTIONS"],
                )["data"]
                if item["target"]["filePath"] == moved_entity_path
            )
            assert moved_target_relation["relationId"] == relation_id
            peer_after_target_move = _semantic_relations(
                client,
                kb_code=kb_code,
                file_path=second_source_path,
                direction="OUTGOING",
                relation_codes=["MENTIONS"],
            )["data"][0]
            assert peer_after_target_move["relationId"] == peer_relation_id
            assert peer_after_target_move["target"]["filePath"] == moved_entity_path

            moved_source_path = f"/moved/anchor-{token}.md"
            _move_file(
                client,
                kb_code=kb_code,
                source_path=source_path,
                target_file_path=moved_source_path,
            )
            moved_source_relation = next(
                item
                for item in _semantic_relations(
                    client,
                    kb_code=kb_code,
                    file_path=moved_source_path,
                    direction="OUTGOING",
                    relation_codes=["MENTIONS"],
                )["data"]
                if item["target"]["filePath"] == moved_entity_path
            )
            assert moved_source_relation["relationId"] == relation_id
            assert moved_source_relation["source"]["filePath"] == moved_source_path

            _update_markdown(
                client,
                kb_code=kb_code,
                file_path=moved_source_path,
                content=(
                    f"# Updated source\n\n{entity_alias} remains linked as "
                    f"[{entity_name}]({moved_entity_path}).\n"
                ),
            )
            rebuilt_a = _relation_assertion_rows(
                kb_code=kb_code, source_path=moved_source_path
            )
            assert [row["discovered_by"] for row in rebuilt_a] == ["MARKDOWN_PARSER"]
            peer_after_a_update = _relation_assertion_rows(
                kb_code=kb_code, source_path=second_source_path
            )
            assert len(peer_after_a_update) == 1
            assert peer_after_a_update[0]["discovered_by"] == "MARKDOWN_PARSER"

            # KE-R3: after the ordinary update removed stale generated output,
            # a forced Discovery replaces only its own producer scope and keeps
            # the Markdown assertion alongside it.
            _build_markdown_index(client, kb_code=kb_code, file_path=moved_source_path)
            rediscovery = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityDiscovery",
                    json={
                        "knCode": kb_code,
                        "filePath": moved_source_path,
                        "maxEntities": 1,
                        "force": True,
                    },
                )
            )
            rediscovery_tasks = _wait_for_batch(
                client,
                kb_code=kb_code,
                batch_id=rediscovery["batchId"],
                expected_task_ids=[rediscovery["tasks"][0]["taskId"]],
            )
            assert rediscovery_tasks[0]["status"] == "SUCCEEDED"
            assertions_after_rediscovery = _relation_assertion_rows(
                kb_code=kb_code, source_path=moved_source_path
            )
            assert sorted(
                row["discovered_by"] for row in assertions_after_rediscovery
            ) == ["ENTITY_DISCOVERY", "MARKDOWN_PARSER"]
            rediscovered_relation = next(
                item
                for item in _semantic_relations(
                    client,
                    kb_code=kb_code,
                    file_path=moved_source_path,
                    direction="OUTGOING",
                    relation_codes=["MENTIONS"],
                )["data"]
                if item["target"]["filePath"] == moved_entity_path
            )
            assert rediscovered_relation["relationId"] == relation_id
            assert rediscovered_relation["assertionCount"] == 2

            # KE-R8/R9: deleting the target keeps both lightweight locators and
            # unbinds the stable target ID. Reimport resolves both producers;
            # deleting source A then removes it from target-facing read views.
            _delete_file(client, kb_code=kb_code, file_path=moved_entity_path)
            assertions_while_deleted = _relation_assertion_rows(
                kb_code=kb_code, source_path=moved_source_path
            )
            assert len(assertions_while_deleted) == 2
            assert all(
                row["target_fs_entry_id"] is None for row in assertions_while_deleted
            )
            assert {row["target_locator_type"] for row in assertions_while_deleted} == {
                "ENTITY_SURFACE",
                "KB_PATH",
            }

            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=moved_entity_path,
                content=_entity_markdown(
                    entity_name=entity_name,
                    aliases=[entity_alias],
                ),
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=moved_entity_path)
            rebound_assertions = _relation_assertion_rows(
                kb_code=kb_code, source_path=moved_source_path
            )
            assert len(rebound_assertions) == 2
            assert all(
                row["target_fs_entry_id"] is not None for row in rebound_assertions
            )
            rebound_relation = next(
                item
                for item in _semantic_relations(
                    client,
                    kb_code=kb_code,
                    file_path=moved_source_path,
                    direction="OUTGOING",
                    relation_codes=["MENTIONS"],
                )["data"]
                if item["target"]["filePath"] == moved_entity_path
            )
            assert rebound_relation["assertionCount"] == 2

            _delete_file(client, kb_code=kb_code, file_path=moved_source_path)
            target_incoming_after_source_delete = _semantic_relations(
                client,
                kb_code=kb_code,
                file_path=moved_entity_path,
                direction="INCOMING",
                relation_codes=["MENTIONS"],
            )
            assert all(
                item["source"]["filePath"] != moved_source_path
                for item in target_incoming_after_source_delete["data"]
            )
            compatibility_incoming = _reference_result(
                client,
                kb_code=kb_code,
                file_path=moved_entity_path,
                direction="inbound",
            )
            assert all(
                item.get("sourcePath") != moved_source_path
                and item.get("filePath") != moved_source_path
                for item in compatibility_incoming["inbound"]
            )
        finally:
            _delete_kb(client, kb_code)
            _delete_kb(client, foreign_kb_code)


@pytest.mark.integration
def test_knowledge_entity_real_whole_enrich_refresh_and_stale_write() -> None:
    """Cover whole-KB Enrich, evidence refresh, and stale-write isolation."""

    _assert_real_runtime_configuration()
    token = uuid4().hex[:10]
    entity_name = f"RefreshEntity{token}"
    entity_alias = f"RefreshAlias{token}"
    entity_path = f"/KnowledgeEntity/refresh-{token}.md"
    no_evidence_path = f"/KnowledgeEntity/no-evidence-{token}.md"
    unsupported_path = f"/KnowledgeEntity/unsupported-{token}.txt"
    evidence_path = f"/evidence/refresh-{token}.md"
    foreign_evidence_path = f"/evidence/foreign-{token}.md"
    local_marker_v1 = f"LOCAL-EVIDENCE-V1-{token}"
    local_marker_v2 = f"LOCAL-EVIDENCE-V2-{token}"
    local_marker_v3 = f"LOCAL-EVIDENCE-V3-{token}"
    foreign_marker = f"FOREIGN-EVIDENCE-{token}"

    with _real_test_client() as client:
        kb_code = _create_kb(client)
        foreign_kb_code = _create_kb(client)
        try:
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                content=_entity_markdown(
                    entity_name=entity_name,
                    aliases=[entity_alias],
                ),
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=entity_path)
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=no_evidence_path,
                content=_entity_markdown(entity_name=f"NoEvidence{token}"),
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=no_evidence_path)
            subject_file_id = int(
                _db_rows(
                    """
                    SELECT kid
                    FROM knowledge_fs_entry
                    WHERE knowledge_base_id = %(knowledge_base_id)s
                      AND virtual_path = %(file_path)s
                      AND is_deleted = FALSE
                    """,
                    {
                        "knowledge_base_id": int(kb_code),
                        "file_path": no_evidence_path,
                    },
                )[0]["kid"]
            )
            _update_metadata(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                operations=[
                    {
                        "propertyName": "subjectFileId",
                        "operation": "set",
                        "valueType": "number",
                        "value": subject_file_id,
                    }
                ],
            )
            stored_subject = _db_rows(
                """
                SELECT metadata.value_number
                FROM knowledge_file_metadata_value metadata
                JOIN knowledge_fs_entry entry ON entry.kid = metadata.fs_entry_id
                WHERE entry.knowledge_base_id = %(knowledge_base_id)s
                  AND entry.virtual_path = %(file_path)s
                  AND metadata.property_name = 'subjectFileId'
                  AND metadata.is_deleted = FALSE
                """,
                {
                    "knowledge_base_id": int(kb_code),
                    "file_path": entity_path,
                },
            )[0]["value_number"]
            assert isinstance(stored_subject, Decimal)
            assert stored_subject == subject_file_id
            _upload_file(
                client,
                kb_code=kb_code,
                file_path=unsupported_path,
                content=b"unsupported entity content",
                content_type="text/plain",
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=unsupported_path)

            evidence_v1 = f"""# Local evidence

{local_marker_v1}. {entity_name}, also known as {entity_alias}, is the stable
integration subject described by this current-KB source.

[{entity_name}]({entity_path})
"""
            _upload_markdown(
                client,
                kb_code=kb_code,
                file_path=evidence_path,
                content=evidence_v1,
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=evidence_path)

            # OpenGauss returns EAV NUMERIC as Decimal. Entity Enrich
            # fingerprinting must canonicalize the subject identity before JSON
            # serialization while preserving the real TIMESTAMPTZ relation version.
            subject_local_eligibility = _eligibility(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                capability="entityEnrich",
            )
            assert subject_local_eligibility["eligibility"] == "ELIGIBLE_AND_STALE"
            assert subject_local_eligibility["reasonCode"] == "NEVER_PROCESSED"

            _upload_markdown(
                client,
                kb_code=foreign_kb_code,
                file_path=foreign_evidence_path,
                content=(
                    f"# Foreign evidence\n\n{foreign_marker}. {entity_name} is "
                    "described only in the foreign knowledge base.\n"
                ),
            )
            _build_markdown_index(
                client,
                kb_code=foreign_kb_code,
                file_path=foreign_evidence_path,
            )

            # KE-E2/E5: whole-KB enumeration selects only the one eligible
            # Markdown entity. Semantic retrieval is fixed to the current KB;
            # the foreign marker must never enter the LLM result.
            whole_batch = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityEnrich",
                    json={
                        "knCode": kb_code,
                        "topK": 5,
                    },
                )
            )
            assert whole_batch["scope"] == "WHOLE_KB"
            assert whole_batch["eligibleCount"] == 1
            assert whole_batch["acceptedCount"] == 1
            assert whole_batch["reusedCount"] == 0
            assert whole_batch["skippedCount"] == 2
            assert len(whole_batch["tasks"]) == 3
            accepted_tasks = [
                item for item in whole_batch["tasks"] if item["status"] == "PENDING"
            ]
            assert len(accepted_tasks) == 1
            assert accepted_tasks[0]["filePath"] == entity_path
            assert {item["filePath"] for item in whole_batch["tasks"]} == {
                entity_path,
                no_evidence_path,
                unsupported_path,
            }
            first_tasks = _wait_for_batch(
                client,
                kb_code=kb_code,
                batch_id=whole_batch["batchId"],
                expected_task_ids=[accepted_tasks[0]["taskId"]],
            )
            first_task = first_tasks[0]
            assert first_task["result"]["evidenceFragmentCount"] >= 1
            first_markdown = _read_markdown(
                client, kb_code=kb_code, file_path=entity_path
            )
            assert first_markdown.strip()
            assert (
                f"{entity_name} is a stable integration-test entity." in first_markdown
            )
            assert foreign_marker not in first_markdown
            first_timeline = _latest_timeline(kb_code=kb_code, file_path=entity_path)
            assert first_timeline is not None
            first_outgoing = _relation_assertion_rows(
                kb_code=kb_code, source_path=entity_path
            )
            assert all(
                row["producer_run_id"] == f"entity-enrich:{first_task['taskId']}"
                for row in first_outgoing
            )
            # KE-E6/M10/R4: rebuilding the changed related source creates a
            # newer incoming relation assertion and makes the entity stale.
            # The second Enrich writes a new timeline entry and replaces every
            # outgoing assertion with the second producer run.
            evidence_v2 = evidence_v1.replace(local_marker_v1, local_marker_v2)
            _update_markdown(
                client,
                kb_code=kb_code,
                file_path=evidence_path,
                content=evidence_v2,
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=evidence_path)
            stale_after_evidence = _eligibility(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                capability="entityEnrich",
            )
            assert stale_after_evidence["eligibility"] == "ELIGIBLE_AND_STALE"
            assert stale_after_evidence["reasonCode"] == "NEW_RELATION"

            second_batch = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityEnrich",
                    json={
                        "knCode": kb_code,
                        "filePath": entity_path,
                        "topK": 5,
                    },
                )
            )
            assert second_batch["acceptedCount"] == 1
            second_tasks = _wait_for_batch(
                client,
                kb_code=kb_code,
                batch_id=second_batch["batchId"],
                expected_task_ids=[second_batch["tasks"][0]["taskId"]],
            )
            second_task = second_tasks[0]
            assert second_task["taskId"] != first_task["taskId"]
            second_markdown = _read_markdown(
                client, kb_code=kb_code, file_path=entity_path
            )
            assert (
                f"{entity_name} is a stable integration-test entity." in second_markdown
            )
            assert "## 集成测试更新" in second_markdown
            second_timeline = _latest_timeline(kb_code=kb_code, file_path=entity_path)
            assert second_timeline is not None
            assert int(second_timeline["kid"]) > int(first_timeline["kid"])
            assert second_timeline["old_checksum"] == first_timeline["new_checksum"]
            second_outgoing = _relation_assertion_rows(
                kb_code=kb_code, source_path=entity_path
            )
            assert all(
                row["producer_run_id"] == f"entity-enrich:{second_task['taskId']}"
                for row in second_outgoing
            )
            assert not {row["kid"] for row in first_outgoing} & {
                row["kid"] for row in second_outgoing
            }

            fresh_after_second = _eligibility(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                capability="entityEnrich",
            )
            assert fresh_after_second["eligibility"] == "ELIGIBLE_BUT_FRESH"
            assert fresh_after_second["reasonCode"] == "NO_NEW_RELATIONS"

            # KE-E7: mutate the target while the Enrich model double is running.
            # The stale referSignature rejects the worker write; the concurrent
            # user's object and metadata remain authoritative. The service
            # currently reports PROCESSING_FAILED rather than the planned
            # STALE_WRITE code, which remains recorded as a matrix gap.
            evidence_v3 = evidence_v2.replace(local_marker_v2, local_marker_v3)
            _update_markdown(
                client,
                kb_code=kb_code,
                file_path=evidence_path,
                content=evidence_v3,
            )
            _build_markdown_index(client, kb_code=kb_code, file_path=evidence_path)
            concurrent_batch = _assert_success(
                client.post(
                    "/api/v1/knowledgeItems/entityEnrich",
                    json={
                        "knCode": kb_code,
                        "filePath": entity_path,
                        "topK": 5,
                    },
                )
            )
            assert concurrent_batch["acceptedCount"] == 1
            concurrent_body = _entity_markdown(
                entity_name=entity_name,
                aliases=[entity_alias],
                body=f"# {entity_name}\n\nCONCURRENT-USER-UPDATE-{token}\n",
            )
            _update_markdown(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                content=concurrent_body,
            )
            concurrent_metadata = _metadata(
                client,
                kb_code=kb_code,
                file_path=entity_path,
                field_names=["fileSignature", "entityName"],
            )
            concurrent_bytes = _download_file(
                client, kb_code=kb_code, file_path=entity_path
            )
            failed_tasks, _ = _wait_for_batch_snapshots(
                client,
                kb_code=kb_code,
                batch_id=concurrent_batch["batchId"],
                expected_task_ids=[concurrent_batch["tasks"][0]["taskId"]],
                require_success=False,
            )
            failed_task = failed_tasks[0]
            assert failed_task["status"] == "FAILED"
            assert failed_task["error"]["errorCode"] == "PROCESSING_FAILED"
            assert "file signature mismatch" in failed_task["error"]["message"]
            assert (
                _download_file(client, kb_code=kb_code, file_path=entity_path)
                == concurrent_bytes
            )
            assert (
                _metadata(
                    client,
                    kb_code=kb_code,
                    file_path=entity_path,
                    field_names=[
                        "fileSignature",
                        "entityName",
                    ],
                )
                == concurrent_metadata
            )
            latest_after_stale = _latest_timeline(
                kb_code=kb_code, file_path=entity_path
            )
            assert latest_after_stale is not None
            assert (
                latest_after_stale["new_checksum"]
                == concurrent_metadata["fileSignature"]["value"]
            )
            assert not _relation_assertion_rows(
                kb_code=kb_code, source_path=entity_path
            )
        finally:
            _delete_kb(client, kb_code)
            _delete_kb(client, foreign_kb_code)


@pytest.mark.integration
def test_knowledge_entity_suffixless_text_mime_fallback_contract() -> None:
    """KE-M7 contract: suffixless text/plain must pass the format gate."""

    _assert_real_runtime_configuration()
    with _real_test_client() as client:
        kb_code = _create_kb(client)
        try:
            file_path = "/formats/no-suffix-text"
            _upload_file(
                client,
                kb_code=kb_code,
                file_path=file_path,
                content=b"suffixless text",
                content_type="Text/Plain; charset=utf-8",
            )
            result = _eligibility(
                client,
                kb_code=kb_code,
                file_path=file_path,
                capability="entityDiscovery",
            )
            assert result["reasonCode"] == "CONTENT_NOT_READY"
        finally:
            _delete_kb(client, kb_code)


@pytest.mark.integration
def test_knowledge_entity_document_kind_backfill_is_real_and_idempotent() -> None:
    """KE-M4: execute SQL 031 twice against real OpenGauss historical rows."""

    _assert_real_runtime_configuration()
    with _real_test_client() as client:
        kb_code = _create_kb(client)
        try:
            paths = (
                "/legacy/original.md",
                "/KnowledgeEntity/legacy-entity.md",
                "/legacy/explicit.md",
                "/legacy/deleted.md",
            )
            for path in paths:
                _upload_markdown(
                    client,
                    kb_code=kb_code,
                    file_path=path,
                    content=f"# {path}\n",
                )
            _update_metadata(
                client,
                kb_code=kb_code,
                file_path="/legacy/explicit.md",
                operations=[
                    {
                        "propertyName": "documentKind",
                        "operation": "set",
                        "valueType": "string",
                        "value": "explicitCustomKind",
                    }
                ],
            )
            _delete_file(client, kb_code=kb_code, file_path="/legacy/deleted.md")

            # The migration test is the one explicitly allowed place where an
            # upgrade-before row is constructed directly in real OpenGauss.
            _db_run(
                """
                DELETE FROM knowledge_file_metadata_value
                WHERE knowledge_base_id = %(knowledge_base_id)s
                  AND property_name = 'documentKind'
                  AND fs_entry_id IN (
                      SELECT kid FROM knowledge_fs_entry
                      WHERE knowledge_base_id = %(knowledge_base_id)s
                        AND virtual_path IN (
                            '/legacy/original.md',
                            '/KnowledgeEntity/legacy-entity.md'
                        )
                  )
                """,
                {"knowledge_base_id": int(kb_code)},
            )
            before = _db_rows(
                """
                SELECT COUNT(*) AS total
                FROM knowledge_file_metadata_value
                WHERE knowledge_base_id = %(knowledge_base_id)s
                  AND property_name = 'documentKind'
                  AND is_deleted = FALSE
                """,
                {"knowledge_base_id": int(kb_code)},
            )[0]["total"]

            project_root = Path(__file__).resolve().parents[3]
            migration_sql = (
                project_root
                / "src/by_qa/knowledge_base/sql/031_knowledge_file_document_kind_backfill.sql"
            ).read_text(encoding="utf-8")
            statements = split_sql_statements(migration_sql)
            _db_run(statements)

            rows_after_first = _db_rows(
                """
                SELECT entry.virtual_path, entry.is_deleted,
                       metadata.value_string AS document_kind
                FROM knowledge_fs_entry entry
                LEFT JOIN knowledge_file_metadata_value metadata
                  ON metadata.fs_entry_id = entry.kid
                 AND metadata.property_name = 'documentKind'
                 AND metadata.is_deleted = FALSE
                WHERE entry.knowledge_base_id = %(knowledge_base_id)s
                  AND entry.entry_type = 'FILE'
                ORDER BY entry.virtual_path
                """,
                {"knowledge_base_id": int(kb_code)},
            )
            by_path = {row["virtual_path"]: row for row in rows_after_first}
            assert by_path["/legacy/original.md"]["document_kind"] == "original"
            assert (
                by_path["/KnowledgeEntity/legacy-entity.md"]["document_kind"]
                == "knowledgeEntity"
            )
            assert (
                by_path["/legacy/explicit.md"]["document_kind"] == "explicitCustomKind"
            )
            assert by_path["/legacy/deleted.md"]["is_deleted"] is True
            assert by_path["/legacy/deleted.md"]["document_kind"] is None

            after_first = _db_rows(
                """
                SELECT COUNT(*) AS total
                FROM knowledge_file_metadata_value
                WHERE knowledge_base_id = %(knowledge_base_id)s
                  AND property_name = 'documentKind'
                  AND is_deleted = FALSE
                """,
                {"knowledge_base_id": int(kb_code)},
            )[0]["total"]
            assert int(after_first) == int(before) + 2

            _db_run(statements)
            after_second = _db_rows(
                """
                SELECT COUNT(*) AS total
                FROM knowledge_file_metadata_value
                WHERE knowledge_base_id = %(knowledge_base_id)s
                  AND property_name = 'documentKind'
                  AND is_deleted = FALSE
                """,
                {"knowledge_base_id": int(kb_code)},
            )[0]["total"]
            assert int(after_second) == int(after_first)
        finally:
            _delete_kb(client, kb_code)
